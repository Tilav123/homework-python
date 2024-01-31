# Задание №1 EXCEL
try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, Border, Side
    file_names = ["1111.xlsx", "2222.xlsx", "3333.xlsx"]
    wb = Workbook()
    ws = wb.active
    ws['A1'] = 1111
    wb.save(file_names[0])
    wb1 = Workbook()
    ws1 = wb1.active
    ws1['A1'] = 2222
    wb1.save(file_names[1])
    wb2 = Workbook()
    ws2 = wb2.active
    ws2['A1'] = 3333
    wb2.save(file_names[2])
    data = []
    for file_name in file_names:
        wb = openpyxl.load_workbook(file_name)
        sheet = wb.active
        file_data = []
        for row in sheet.iter_rows(values_only=True):
            file_data.append(row)
        data.extend(file_data)
    data.sort(reverse=True)
    output_wb = openpyxl.Workbook()
    output_sheet = output_wb.active
    for row_idx, row in enumerate(data, start=1):
        for col_idx, cell_value in enumerate(row, start=1):
            output_sheet.cell(row=row_idx, column=col_idx, value=cell_value)
            output_sheet.cell(row=row_idx, column=col_idx).font = Font(bold=True)
            border_style = Side(border_style="thin", color="000000")
            border = Border(top=border_style, bottom=border_style, left=border_style, right=border_style)
            output_sheet.cell(row=row_idx, column=col_idx).border = border
    output_wb.save("sorted_data.xlsx")
except:
    print("Ошибка")
# Задание №2 JSON
try:
    import json
    import requests
    import os
    os.makedirs('todos_output')
    response = requests.get('https://jsonplaceholder.typicode.com/todos/')
    todos = response.json()
    with open('todos.json', 'w') as file:
        json.dump(todos, file)
    with open('todos.json', 'r') as file:
        todos_data = json.load(file)
        # Записываю каждый словарь в отдельный json файл
        p=0
        for i in todos_data:
            p+=1
            with open("todos_output/"+str(p)+'.json', 'w') as file:
                json.dump(i, file)
except:
    print("Ошибка")
# Задание №3 WORD
import docx
doc = docx.Document()
doc.add_paragraph('Hello Python')
doc.save('word.docx')
doc2=docx.Document("word.docx")
for paragraph in doc2.paragraphs:
    for run in paragraph.runs:
        run.bold = True
doc2.save("new_word.docx")